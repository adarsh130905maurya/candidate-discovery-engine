"""
data_loader.py — Load & validate raw hackathon data.

============================================================================
WHAT THIS MODULE DOES (read me — I'm written for non-technical teammates)
============================================================================
This is the FIRST module in our pipeline. It does four things:

  1. LOAD CANDIDATES
     Reads data/candidates.jsonl (a 487 MB file with ~100,000 candidate
     records, one JSON object per line) into a pandas DataFrame.
     We STREAM the file line-by-line so we never hold the whole 487 MB
     in RAM at once.

  2. VALIDATE SCHEMA
     Checks every candidate against data/candidate_schema.json using the
     `jsonschema` library. Bad records are flagged (NOT crashed on) so
     the rest of the pipeline keeps running.

  3. EXTRACT DOCX TEXT
     The job description, signals doc, and submission spec are all Word
     documents. We pull the text out of them (including tables) as plain
     strings that other modules can read.

  4. EXPOSE FUNCTIONS
     Other modules import these (this is the "contract" Person 1 owes
     Persons 2, 3, and 4):

         from data_loader import load_candidates
         from data_loader import get_job_description
         from data_loader import get_signals_doc
         from data_loader import get_submission_spec

============================================================================
HOW TO RUN IT STANDALONE (for a quick look at the data)
============================================================================
    python src/data_loader.py            # full summary
    python src/data_loader.py --limit 50 # only scan first 50 lines (fast)

============================================================================
"""

from __future__ import annotations

import argparse          # parse --limit command-line flag
import json              # parse .jsonl lines and the schema file
import sys
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional

import pandas as pd

# --- Optional imports -----------------------------------------------------
# We wrap these in try/except so the module still imports even if a teammate
# hasn't run `pip install -r requirements.txt` yet — they'll get a friendly
# error instead of a cryptic ImportError.
try:
    from docx import Document as _DocxDocument  # python-docx
except ImportError:  # pragma: no cover
    _DocxDocument = None

try:
    import jsonschema  # schema validation
except ImportError:  # pragma: no cover
    jsonschema = None

# --- Our own config -------------------------------------------------------
# `config` lives in the same folder (src/). We add this folder to sys.path
# so the import works whether you run from project root or from src/.
_THIS_DIR = Path(__file__).resolve().parent
if str(_THIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THIS_DIR))

from config import CONFIG, PATHS  # noqa: E402  (path setup above)


# ============================================================================
# 1. DOCX TEXT EXTRACTION
# ============================================================================
def extract_docx_text(docx_path: Path | str) -> str:
    """
    Read a .docx file and return ALL its text as one big string.

    Includes BOTH paragraphs AND tables (tables are joined with ' | '
    between cells, like a mini CSV row). This matters because the
    submission_spec.docx has important rules inside tables.

    Args:
        docx_path: path to a .docx file (str or pathlib.Path)

    Returns:
        A single string with the document's full text.

    Raises:
        RuntimeError: if python-docx isn't installed, or the file is missing.
    """
    docx_path = Path(docx_path)

    # Friendly error if python-docx is missing
    if _DocxDocument is None:
        raise RuntimeError(
            "python-docx is not installed. Run:\n"
            "    pip install python-docx"
        )
    if not docx_path.exists():
        raise FileNotFoundError(
            f"DOCX file not found: {docx_path}\n"
            f"Make sure all hackathon files are inside the data/ folder."
        )

    doc = _DocxDocument(str(docx_path))
    parts: List[str] = []

    # --- 1a. Paragraphs (the normal body text) ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # --- 1b. Tables (joined as "cell | cell | cell" rows) ---
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # skip fully-empty rows
                parts.append(" | ".join(cells))

    return "\n".join(parts)


# --- Cached accessors for the three docx documents -------------------------
# We use a tiny cache dict so we only read each file ONCE even if the
# function is called many times. Reading a .docx is slow-ish.
_DOCX_CACHE: Dict[str, str] = {}


def get_job_description() -> str:
    """Return the job description as plain text (cached after first call)."""
    key = "job_description"
    if key not in _DOCX_CACHE:
        _DOCX_CACHE[key] = extract_docx_text(PATHS["job_description"])
    return _DOCX_CACHE[key]


def get_signals_doc() -> str:
    """Return redrob_signals_doc.docx as plain text (cached)."""
    key = "signals_doc"
    if key not in _DOCX_CACHE:
        _DOCX_CACHE[key] = extract_docx_text(PATHS["signals_doc"])
    return _DOCX_CACHE[key]


def get_submission_spec() -> str:
    """Return submission_spec.docx as plain text (cached)."""
    key = "submission_spec"
    if key not in _DOCX_CACHE:
        _DOCX_CACHE[key] = extract_docx_text(PATHS["submission_spec"])
    return _DOCX_CACHE[key]


def get_readme() -> str:
    """Return README.docx as plain text (cached)."""
    key = "readme"
    if key not in _DOCX_CACHE:
        _DOCX_CACHE[key] = extract_docx_text(PATHS["readme"])
    return _DOCX_CACHE[key]


# ============================================================================
# 2. JSONL STREAMING + SCHEMA VALIDATION
# ============================================================================
def iter_candidate_records(
    limit: Optional[int] = None,
) -> Iterator[Dict[str, Any]]:
    """
    Yield candidate records one at a time from candidates.jsonl.

    WHY A GENERATOR (yield) INSTEAD OF A LIST?
        The file is 487 MB. If we loaded it all into a list we'd use ~2 GB
        of RAM. A generator reads ONE line at a time, so memory stays flat
        no matter how big the file is.

    Args:
        limit: stop after this many records (useful for quick tests).
               None means "read the whole file".

    Yields:
        One parsed candidate dict per line.

    Raises:
        FileNotFoundError: if candidates.jsonl is missing.
    """
    path = PATHS["candidates"]
    if not path.exists():
        raise FileNotFoundError(
            f"Candidates file not found: {path}\n"
            f"Put candidates.jsonl inside the data/ folder."
        )

    count = 0
    # Open in text mode with utf-8 (the file has accented names, emojis etc.)
    with open(path, "r", encoding=CONFIG["candidates_encoding"]) as fh:
        for line_num, raw_line in enumerate(fh, start=1):
            line = raw_line.strip()
            if not line:
                continue  # skip blank lines

            try:
                record = json.loads(line)
            except json.JSONDecodeError as exc:
                # One bad line should NOT kill the whole pipeline.
                # Print a warning and move on.
                print(f"  [warn] line {line_num}: bad JSON ({exc.msg}); skipped",
                      file=sys.stderr)
                continue

            yield record
            count += 1

            if limit is not None and count >= limit:
                return  # stop early — used by --limit flag


def load_candidates(
    limit: Optional[int] = None,
    validate: bool = True,
) -> pd.DataFrame:
    """
    Read candidates.jsonl into a pandas DataFrame.

    Args:
        limit: max number of records to load (None = whole file).
        validate: if True (default), validate each record against the
                  schema and add an `_schema_valid` column.

    Returns:
        A DataFrame with one row per candidate. Top-level keys become
        columns: candidate_id, profile, career_history, education,
        skills, certifications, languages, redrob_signals.
        (Flattening those nested columns into useful columns is the
        job of data_cleaner.py — Person 1's next step.)
    """
    # Load the JSON Schema once (for validation, if requested)
    schema = None
    if validate:
        try:
            with open(PATHS["candidate_schema"], "r", encoding="utf-8") as f:
                schema = json.load(f)
        except FileNotFoundError:
            print("  [warn] candidate_schema.json not found; skipping validation",
                  file=sys.stderr)
            validate = False
        except json.JSONDecodeError:
            print("  [warn] candidate_schema.json is malformed; skipping validation",
                  file=sys.stderr)
            validate = False

    # If jsonschema isn't installed, we can't validate
    if validate and jsonschema is None:
        print("  [warn] jsonschema library not installed; skipping validation.\n"
              "         Run: pip install jsonschema",
              file=sys.stderr)
        validate = False

    # Build a jsonschema Validator once (cheap to reuse)
    validator = None
    if validate:
        # Draft7Validator matches the "$schema": "draft-07" in our schema
        validator = jsonschema.Draft7Validator(schema)

    records: List[Dict[str, Any]] = []
    valid_count = 0
    invalid_count = 0

    for rec in iter_candidate_records(limit=limit):
        if validator is not None:
            errors = list(validator.iter_errors(rec))
            if errors:
                invalid_count += 1
                rec["_schema_valid"] = False
                rec["_schema_errors"] = "; ".join(
                    f"{'.'.join(str(p) for p in e.absolute_path) or '<root>'}: {e.message}"
                    for e in errors[:3]  # keep first 3 errors to stay short
                )
            else:
                valid_count += 1
                rec["_schema_valid"] = True
                rec["_schema_errors"] = ""
        else:
            rec["_schema_valid"] = None
            rec["_schema_errors"] = ""
        records.append(rec)

    df = pd.DataFrame(records)

    # Report
    print(f"[load_candidates] Loaded {len(df):,} candidate records.")
    if validator is not None:
        print(f"  Schema valid:   {valid_count:,}")
        print(f"  Schema invalid: {invalid_count:,}")
    return df


def get_column_summary(df: pd.DataFrame) -> Dict[str, str]:
    """
    Return a quick human-readable summary of each column's type.

    Useful for the exploration doc — tells us which columns are text,
    which are nested dicts/lists, etc.
    """
    summary: Dict[str, str] = {}
    for col in df.columns:
        if col.startswith("_"):
            continue
        # Peek at first non-null value
        sample = df[col].dropna().iloc[0] if len(df[col].dropna()) else None
        if sample is None:
            summary[col] = "empty"
        else:
            summary[col] = type(sample).__name__
    return summary


# ============================================================================
# 3. STANDALONE SUMMARY (python src/data_loader.py)
# ============================================================================
def _main() -> int:
    """Print a readable summary when the file is run directly."""
    parser = argparse.ArgumentParser(description="Load & inspect hackathon data.")
    parser.add_argument(
        "--limit", type=int, default=None,
        help="Only read the first N candidate lines (for a quick test).",
    )
    args = parser.parse_args()

    # Use the debug cap from config if --limit wasn't given
    limit = args.limit if args.limit is not None else CONFIG["max_candidates_debug"]

    print("=" * 72)
    print(" CANDIDATE DISCOVERY ENGINE — data_loader.py summary")
    print("=" * 72)

    # --- Load candidates ---
    print(f"\n[1/4] Loading candidates (limit={limit})...")
    df = load_candidates(limit=limit, validate=True)

    print("\n[2/4] Columns available:")
    for col in df.columns:
        print(f"   - {col}")

    # --- Schema result ---
    if "_schema_valid" in df.columns and df["_schema_valid"].notna().any():
        n_valid = int(df["_schema_valid"].sum())
        print(f"\n[3/4] Schema validation: {n_valid}/{len(df)} records valid.")
        if n_valid < len(df):
            bad = df[df["_schema_valid"] == False].head(3)
            for _, row in bad.iterrows():
                print(f"   - {row['candidate_id']}: {row['_schema_errors']}")
    else:
        print("\n[3/4] Schema validation skipped.")

    # --- Job description preview ---
    print("\n[4/4] Job description preview (first 300 chars):")
    try:
        jd = get_job_description()
        print(jd[:300] + ("..." if len(jd) > 300 else ""))
    except Exception as exc:
        print(f"   ERROR: {exc}")

    # --- Signals doc preview ---
    print("\nSignals doc preview (first 200 chars):")
    try:
        sd = get_signals_doc()
        print(sd[:200] + ("..." if len(sd) > 200 else ""))
    except Exception as exc:
        print(f"   ERROR: {exc}")

    print("\n" + "=" * 72)
    print(" Done. Import this module's functions from other code:")
    print("   from data_loader import load_candidates, get_job_description,")
    print("                                 get_signals_doc, get_submission_spec")
    print("=" * 72)
    return 0


if __name__ == "__main__":
    raise SystemExit(_main())
